import os
import re
import math
import json
import torch
import nltk
from io import BytesIO
from docx import Document
from nltk.corpus import wordnet as wn
from transformers import AutoModelForCausalLM, AutoTokenizer
import requests
import gradio as gr
from latex_parser import mask_latex, unmask_latex

# Rileva se siamo su HF Spaces (ZeroGPU) o su Colab/locale
IS_HF_SPACES = os.environ.get("SPACE_ID") is not None
try:
    import spaces
    if not IS_HF_SPACES:
        raise ImportError("Not on HF Spaces")
except ImportError:
    # Su Colab/locale, creiamo un decoratore finto che non fa nulla
    class _FakeSpaces:
        @staticmethod
        def GPU(duration=60):
            def decorator(fn):
                return fn
            return decorator
    spaces = _FakeSpaces()

# --- SCARICO RISORSE NLTK ---
nltk.download('wordnet', quiet=True)
nltk.download('omw-1.4', quiet=True)

# --- CHIAVI API ---
MISTRAL_API_KEY = os.environ.get("MISTRAL_API_KEY", "")

# --- INIZIALIZZAZIONE MODELLO (CARICATO SU CUDA A LIVELLO GLOBALE PER ZEROGPU) ---
model_id = "Qwen/Qwen3.5-0.8B"
tokenizer = AutoTokenizer.from_pretrained(model_id)
if tokenizer.pad_token is None:
    tokenizer.pad_token = tokenizer.eos_token

try:
    # Cerchiamo di allocarlo su CUDA. In locale potrebbe fallire se non c'è GPU, ma su ZeroGPU c'è sempre.
    model = AutoModelForCausalLM.from_pretrained(model_id, device_map="cuda", trust_remote_code=True)
except Exception:
    model = AutoModelForCausalLM.from_pretrained(model_id, device_map="cpu", trust_remote_code=True)


AI_CLICHES = [
    "delve", "tapestry", "testament", "underscore", "vibrant", "landscape", "pivotal", "showcase", "intricate", "crucial",
    "fostering", "garner", "highlight", "interplay", "emphasizing", "enduring", "enhance", "multifaceted", "synergy", "paradigm",
    "catalyst", "navigate", "realm", "unveil", "cornerstone", "essence", "dynamic", "robust", "meticulous", "nuanced",
    "approfondire", "arazzo", "testimonianza", "sottolineare", "vibrante", "panorama", "paesaggio", "cruciale", "mostrare", "intricato",
    "promuovere", "raccogliere", "interazione", "enfatizzare", "duraturo", "migliorare", "immersione", "immergiamoci",
    "poliedrico", "sinergia", "paradigma", "catalizzatore", "navigare", "regno", "svelare", "fondamento", "essenza", "dinamico",
    "robusto", "meticoloso", "scrupoloso", "sfumato", "tuttavia", "inoltre", "infine", "complessivo", "affascinante", "sorprendente"
]

SYSTEM_PROMPT_HUMANIZER = """
Sei un editor esperto. Il tuo compito è sostituire solo 1 o 2 parole nella frase per aggirare i detector AI.

REGOLE FONDAMENTALI:
1. FEDELTÀ ESTREMA: Non stravolgere o riscrivere l'intera frase. Mantieni la struttura, la punteggiatura e la lunghezza identiche all'originale.
2. SOSTITUZIONE MINIMA: Cambia al massimo 1 o 2 parole (quelle che sembrano più prevedibili, banali o "cliché" da intelligenza artificiale) con sinonimi leggermente meno comuni ma perfettamente naturali per il contesto.
3. MULTILINGUA: La frase può essere in Italiano o in Inglese. Mantieni SEMPRE la lingua originale della frase. Non tradurre per nessun motivo.
4. PRESERVA IL SIGNIFICATO: Il senso originale della frase deve rimanere assolutamente identico.
5. PRESERVA I SEGNAPOSTO: DEVI conservare intatti e nella giusta posizione tutti i placeholder come [MATH_n], [CITE_n], [CMD_n]. Non modificarli per nessun motivo.

OUTPUT:
Restituisci ESCLUSIVAMENTE la frase modificata. Nessuna introduzione, nessuna spiegazione, nessuna virgoletta iniziale o finale.
"""

def apply_algorithmic_rules(text):
    if not text: return text
    text = text.replace("—", ", ").replace("–", ", ")
    text = text.replace(" — ", ", ").replace(" -- ", ", ")
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = text.replace("“", '"').replace("”", '"')
    return text

def split_into_sentences(text):
    # Separiamo per punteggiatura finale o per paragrafi (doppi a capo)
    sentences = re.split(r'(?<=[.!?])\s+|\n\n+', text)
    return [s for s in sentences if s.strip()]

def calculate_burstiness(sentences):
    lengths = [len(s.split()) for s in sentences if s.strip()]
    if not lengths: return 0.0
    mean_len = sum(lengths)/len(lengths)
    variance = sum((l - mean_len)**2 for l in lengths)/len(lengths)
    return round(math.sqrt(variance), 2)

def process_sentence(sentence_text, latex_registry, is_latex):
    """Calcola perplexity della frase e dei singoli token. NON È DECORATO perché viene chiamato da handler decorati."""
    qwen_text = sentence_text
    if is_latex and latex_registry:
        for mask in latex_registry.keys():
            qwen_text = qwen_text.replace(mask, "formula")
            
    clean_qwen = qwen_text.replace("formula", "").strip()
    if is_latex and not clean_qwen:
        words_data = []
        for w in sentence_text.split():
            word_str = w
            is_mask = False
            for mask, real_val in latex_registry.items():
                if mask in word_str:
                    word_str = word_str.replace(mask, real_val)
                    is_mask = True
            words_data.append({
                "word": word_str,
                "isLowPpl": False,
                "isCliche": False,
                "isMask": is_mask
            })
        return {
            "text": sentence_text,
            "words": words_data,
            "isCritical": False,
            "ppl": 1000.0,
            "aiScore": 0
        }

    words_data = []
    encodings = tokenizer(qwen_text, return_tensors="pt", truncation=True, max_length=512)
    # Su ZeroGPU, le operazioni tensor all'interno delle funzioni @spaces.GPU possono usare .to("cuda")
    try:
        input_ids = encodings.input_ids[0].to(model.device)
    except:
        input_ids = encodings.input_ids[0]
    
    try:
        with torch.no_grad():
            outputs = model(input_ids.unsqueeze(0), labels=input_ids.unsqueeze(0))
        logits = outputs.logits[0, :-1, :]
        labels = input_ids[1:]
        loss_fct = torch.nn.CrossEntropyLoss(reduction='none')
        losses = loss_fct(logits, labels).tolist()
        losses = [0.0] + losses
        avg_loss = sum(losses)/len(losses) if losses else 0
    except Exception:
        losses = [10.0] * len(input_ids)
        avg_loss = 10.0
        
    words = sentence_text.split()
    token_idx = 0
    for w in words:
        w_tokens = tokenizer(w, add_special_tokens=False).input_ids
        num_tokens = max(1, len(w_tokens))
        
        w_losses = []
        for _ in range(num_tokens):
            if token_idx < len(losses):
                w_losses.append(losses[token_idx])
                token_idx += 1
                
        word_loss = sum(w_losses) / len(w_losses) if w_losses else 10.0
            
        try:
            ppl = math.exp(word_loss)
        except OverflowError:
            ppl = 1000.0
            
        clean_w = "".join(c for c in w if c.isalpha())
        is_low_ppl = (ppl < 15.0 and len(clean_w) > 3)
        is_cliche = clean_w.lower() in AI_CLICHES
        
        words_data.append({
            "word": w,
            "isLowPpl": is_low_ppl,
            "isCliche": is_cliche
        })
        
    sentence_ppl = round(math.exp(avg_loss) if avg_loss < 100 else 1000.0, 1)
    cliche_count = sum(1 for w in words_data if w["isCliche"])
    is_critical = (sentence_ppl < 30.0) or (cliche_count > 0)
    
    ai_score = round(max(0, min(100, 100 - (sentence_ppl * 2) + (cliche_count * 10))))
    
    if is_latex and latex_registry:
        for w_dict in words_data:
            for mask in latex_registry.keys():
                if mask in w_dict["word"]:
                    w_dict["isMask"] = True
                    w_dict["isLowPpl"] = False
                    w_dict["isCliche"] = False
    
    return {
        "text": sentence_text,
        "words": words_data,
        "isCritical": is_critical,
        "ppl": sentence_ppl,
        "aiScore": ai_score
    }

def get_mistral_synonyms(word, context_sentence="", lang="Italiano"):
    if not MISTRAL_API_KEY:
        return []
    url = "https://api.mistral.ai/v1/chat/completions"
    headers = {
        "Content-Type": "application/json",
        "Accept": "application/json",
        "Authorization": f"Bearer {MISTRAL_API_KEY}"
    }
    
    lang_instruction = "Italiano" if lang == "Italiano" else "English"
    data = {
        "model": "mistral-small-latest",
        "messages": [
            {"role": "system", "content": f"Sei un dizionario dei sinonimi. Restituisci SOLO una lista di 5 sinonimi appropriati nella lingua {lang_instruction}. I sinonimi devono essere separati da virgola. Nessuna introduzione, nessuna spiegazione."},
            {"role": "user", "content": f"Fornisci 5 sinonimi per la parola '{word}' nel contesto di questa frase: '{context_sentence}'"}
        ]
    }
    try:
        response = requests.post(url, headers=headers, json=data)
        response.raise_for_status()
        raw_output = response.json()["choices"][0]["message"]["content"]
        return [s.strip() for s in raw_output.split(",") if s.strip()]
    except:
        return []

def get_offline_synonyms(word, context_sentence="", lang="Italiano"):
    clean_word = "".join(c for c in word if c.isalpha()).lower()
    syns = set()
    
    wn_lang = 'ita' if lang == 'Italiano' else 'eng'
    
    synsets = wn.synsets(clean_word, lang=wn_lang)
    for syn in synsets:
        for lemma in syn.lemma_names(wn_lang):
            if lemma.lower() != clean_word:
                syns.add(lemma.replace('_', ' '))
                
    syns_list = list(syns)[:5]
    if not syns_list and context_sentence:
        syns_list = get_mistral_synonyms(clean_word, context_sentence, lang)
    return syns_list


def calculate_synonym_scores(sentence_data, word_idx, synonyms):
    results = []
    if not synonyms:
        return results
        
    left_context = " ".join([w["word"] for w in sentence_data["words"][:word_idx]])
    right_context = " ".join([w["word"] for w in sentence_data["words"][word_idx+1:]])
    
    texts = []
    for syn in synonyms:
        new_text = f"{left_context} {syn} {right_context}".strip()
        texts.append(new_text)
        
    encodings = tokenizer(texts, return_tensors="pt", padding=True, truncation=True, max_length=512)
    
    try:
        with torch.no_grad():
            ids = encodings.input_ids.to(model.device)
            attention_mask = encodings.attention_mask.to(model.device)
            outputs = model(ids, attention_mask=attention_mask, labels=ids)
            
            logits = outputs.logits
            shift_logits = logits[..., :-1, :].contiguous()
            shift_labels = ids[..., 1:].contiguous()
            shift_mask = attention_mask[..., 1:].contiguous()
            
            loss_fct = torch.nn.CrossEntropyLoss(reduction='none')
            loss = loss_fct(shift_logits.view(-1, shift_logits.size(-1)), shift_labels.view(-1))
            loss = loss.view(shift_labels.size())
            
            loss = loss * shift_mask
            seq_lengths = shift_mask.sum(dim=1)
            seq_lengths = torch.clamp(seq_lengths, min=1)
            seq_losses = loss.sum(dim=1) / seq_lengths
            
            for i, syn in enumerate(synonyms):
                try:
                    score = math.exp(seq_losses[i].item())
                except OverflowError:
                    score = 1000.0
                results.append({"word": syn, "score": score})
    except Exception as e:
        print(f"[DEBUG] Errore nel calcolo batched dei sinonimi: {e}")
        # Fallback in caso di errore
        for syn in synonyms:
            results.append({"word": syn, "score": 1000.0})
            
    results.sort(key=lambda x: x["score"], reverse=True)
    return results

def rewrite_with_mistral(text, sentence_data=None, temperature=0.7, lang="Italiano", context_before="", context_after=""):
    if not MISTRAL_API_KEY:
        return text
        
    lang_prompt = "italiano" if lang == "Italiano" else "inglese"
    prompt = f"Modifica minimamente la seguente frase in {lang_prompt} sostituendo al massimo 1 o 2 parole con sinonimi perfetti per il contesto, mantenendo intatto tutto il resto. NON aggiungere formattazioni markdown (niente asterischi o grassetti).\n\n"
    
    if context_before or context_after:
        prompt += f"CONTESTO PRECEDENTE: {context_before}\n" if context_before else ""
        prompt += f"FRASE DA MODIFICARE: {text}\n"
        prompt += f"CONTESTO SUCCESSIVO: {context_after}\n\n"
        prompt += "Assicurati che la frase riscritta si colleghi in modo fluido al contesto precedente e successivo e restituisci SOLO la FRASE DA MODIFICARE, senza aggiungere altro."
    else:
        prompt += f"Testo originale:\n\n{text}"
    
    if sentence_data:
        suggestions = []
        for i, w_dict in enumerate(sentence_data["words"]):
            if w_dict.get("isLowPpl") or w_dict.get("isCliche"):
                syns = get_offline_synonyms(w_dict["word"], text, lang)
                if syns:
                    scores = calculate_synonym_scores(sentence_data, i, syns)
                    if scores:
                        top_syns = ", ".join([f"{s['word']}" for s in scores[:3]])
                        suggestions.append(f"- {w_dict['word']}: {top_syns}")
        if suggestions:
            prompt += "\n\nEcco alcuni sinonimi raccomandati per le parole da sostituire. SCEGLINE AL MASSIMO UNO O DUE (i più naturali per il contesto) e usali al posto delle parole originali. IGNORA tutti gli altri suggerimenti e non stravolgere la frase:\n" + "\n".join(suggestions)
            
    url = "https://api.mistral.ai/v1/chat/completions"
    headers = {
        "Content-Type": "application/json",
        "Accept": "application/json",
        "Authorization": f"Bearer {MISTRAL_API_KEY}"
    }
    data = {
        "model": "mistral-small-latest",
        "temperature": temperature,
        "messages": [
            {"role": "system", "content": SYSTEM_PROMPT_HUMANIZER},
            {"role": "user", "content": prompt}
        ]
    }
    try:
        response = requests.post(url, headers=headers, json=data)
        response.raise_for_status()
        return response.json()["choices"][0]["message"]["content"].strip()
    except:
        return text

# --- FUNZIONI GRADIO ---

PAGE_SIZE = 20

def build_html(processed_sentences, valid_sentences=None, synonyms_payload=None, current_page=1, page_size=PAGE_SIZE):
    import html
    import json
    
    processed_count = sum(1 for s in processed_sentences if s is not None)
    global_ai_score = 0
    if processed_count > 0:
        global_ai_score = sum(s["aiScore"] for s in processed_sentences if s is not None) / processed_count
        
    total_sentences = len(valid_sentences) if valid_sentences else len(processed_sentences)
    progress_pct = (processed_count / total_sentences) * 100 if total_sentences > 0 else 100
    
    start_idx = (current_page - 1) * page_size
    end_idx = min(start_idx + page_size, total_sentences)
    
    page_sentences_data = []
    pending_texts = []
    
    if valid_sentences:
        for i in range(start_idx, end_idx):
            if processed_sentences[i] is not None:
                page_sentences_data.append(processed_sentences[i])
            else:
                pending_texts.append(valid_sentences[i])
    else:
        for i in range(start_idx, min(start_idx + page_size, len(processed_sentences))):
            if processed_sentences[i] is not None:
                page_sentences_data.append(processed_sentences[i])

    sentences_json = json.dumps(page_sentences_data).replace("'", "\\'")
    synonyms_json = json.dumps(synonyms_payload) if synonyms_payload else "null"
    pending_json = json.dumps(pending_texts).replace("'", "\\'")
    
    inner_html = f"""
    <html>
    <head>
    <style>
      body {{ font-family: 'Inter', sans-serif; background-color: transparent; color: #f8fafc; padding: 20px; line-height: 1.9; font-size: 16px; margin: 0; overflow-y: auto; }}
      .progress-bar {{ background: #334155; height: 8px; border-radius: 4px; margin-bottom: 20px; overflow: hidden; }}
      .progress-fill {{ background: #3b82f6; height: 100%; transition: width 0.3s; }}
      .sentence {{ 
          display: block; padding: 12px 18px; margin-bottom: 12px; 
          background: #1e293b; border-radius: 8px; border: 1px solid rgba(255,255,255,0.05);
          transition: transform 0.2s, box-shadow 0.2s, border-color 0.2s;
          position: relative; line-height: 1.6; color: #cbd5e1;
      }}
      .sentence.pending {{
          opacity: 0.4;
          filter: grayscale(100%);
          pointer-events: none;
      }}
      .sentence.critical {{ background-color: rgba(248, 113, 113, 0.05); border-left: 4px solid #ef4444; }}
      .sentence.critical:hover {{ background-color: rgba(248, 113, 113, 0.1); }}
      .sentence:hover .rewrite-btn {{ display: inline-block; }}
      
      .rewrite-btn {{
          display: none; position: absolute; top: 16px; right: 16px; 
          background: linear-gradient(135deg, #3b82f6, #8b5cf6); color: white;
          font-size: 12px; padding: 6px 12px; border-radius: 12px; cursor: pointer; 
          box-shadow: 0 4px 10px rgba(0,0,0,0.5); white-space: nowrap; font-weight: bold; 
          z-index: 100; user-select: none;
      }}
      .sentence.critical .rewrite-btn {{ background: linear-gradient(135deg, #ef4444, #f59e0b); }}
      .rewrite-btn:hover {{ transform: translateY(-2px); filter: brightness(1.1); }}
      
      .metrics-badge {{
          font-size: 11px; color: rgba(255,255,255,0.5); background: rgba(0,0,0,0.2);
          padding: 2px 6px; border-radius: 4px; margin-left: 6px; vertical-align: middle; user-select: none;
      }}
      .sentence.critical .metrics-badge {{ color: #f87171; background: rgba(248, 113, 113, 0.1); border: 1px solid rgba(248, 113, 113, 0.3); }}
  
      .word {{ cursor: pointer; transition: color 0.2s; padding: 0 1px; display: inline-block; }}
      .word.low-ppl {{ color: #f87171; font-weight: 600; }}
      .word.cliche {{ color: #eab308; font-weight: 600; }}
      .word.mask {{ background: #334155; color: #94a3b8; padding: 0 4px; border-radius: 4px; font-family: monospace; font-size: 0.85em; cursor: default; }}
      .word:hover {{ background-color: rgba(255,255,255,0.1); border-radius: 3px; }}
      
      #context-menu {{
        display: none; position: fixed; background: #1e293b; border: 1px solid rgba(255,255,255,0.1);
        border-radius: 8px; box-shadow: 0 4px 15px rgba(0,0,0,0.5); z-index: 9999; padding: 5px 0; min-width: 150px;
      }}
      .menu-item {{ padding: 8px 15px; cursor: pointer; color: #e2e8f0; font-size: 14px; }}
      .menu-item:hover {{ background: #3b82f6; color: white; }}
      .menu-item .score {{ float: right; font-size: 11px; color: #94a3b8; margin-left: 10px; }}
      .menu-item:hover .score {{ color: #e0f2fe; }}
      .loader {{ padding: 8px 15px; font-size: 12px; color: #94a3b8; font-style: italic; }}
    </style>
    </head>
    <body>
    <div class="progress-bar"><div class="progress-fill" style="width: {progress_pct}%"></div></div>
        <div style="display: flex; justify-content: space-between; margin-bottom: 8px;">
            <span style="font-weight: bold; color: #e2e8f0;">Punteggio Globale AI: {int(global_ai_score)}%</span>
            <span style="color: #94a3b8; font-size: 12px;">{processed_count} / {total_sentences} frasi totali</span>
        </div>
    <div class="text-container" id="text-container"></div>
    <div id="context-menu"></div>
    
    <script>
      function sendToGradio(payload) {{
          window.parent.postMessage({{ type: "gradio_action", payload: payload }}, "*");
      }}
      
      function requestResize(extraHeight = 50) {{
          window.parent.postMessage({{ type: "resize_iframe", height: document.body.scrollHeight + extraHeight }}, "*");
      }}
  
      let currentSIdx = null;
      let currentWIdx = null;
      let sentencesData = {sentences_json};
      let pendingTexts = {pending_json};
      let synonymsPayload = {synonyms_json};
      let pageStartIdx = {start_idx};
      
      function renderText() {{
          const container = document.getElementById("text-container");
          container.innerHTML = "";
          
          sentencesData.forEach((s, sIdx) => {{
              let sSpan = document.createElement("span");
              sSpan.className = "sentence" + (s.isCritical ? " critical" : "");
              
              let btn = document.createElement("div");
              btn.className = "rewrite-btn";
              btn.textContent = s.isCritical ? "⚠️ Riscrivi AI" : "🪄 Riscrivi AI";
              btn.onclick = (e) => {{
                  e.stopPropagation();
                  btn.textContent = "⏳ Riscrittura in corso...";
                  btn.style.pointerEvents = "none";
                  sendToGradio({{ action: "rewrite_sentence", sentence_idx: sIdx + pageStartIdx, text: s.text }});
              }};
              sSpan.appendChild(btn);
              
              let btn_edit = document.createElement("div");
              btn_edit.className = "rewrite-btn";
              btn_edit.style.backgroundColor = "#475569";
              btn_edit.style.right = "130px";
              btn_edit.textContent = "✏️ Modifica";
              btn_edit.onclick = (e) => {{
                  e.stopPropagation();
                  let manualText = prompt("Modifica manualmente la frase:", s.text);
                  if (manualText !== null && manualText.trim() !== "" && manualText !== s.text) {{
                      btn_edit.textContent = "⏳ Calcolo...";
                      btn_edit.style.pointerEvents = "none";
                      sendToGradio({{ action: "manual_edit", sentence_idx: sIdx + pageStartIdx, text: manualText.trim() }});
                  }}
              }};
              sSpan.appendChild(btn_edit);
              
              if (s.original_text && s.original_text !== s.text) {{
                  let btn_undo = document.createElement("div");
                  btn_undo.className = "rewrite-btn";
                  btn_undo.style.backgroundColor = "#ef4444";
                  btn_undo.style.right = "240px";
                  btn_undo.textContent = "↩️ Annulla";
                  btn_undo.onclick = (e) => {{
                      e.stopPropagation();
                      btn_undo.textContent = "⏳ Ripristino...";
                      btn_undo.style.pointerEvents = "none";
                      sendToGradio({{ action: "undo", sentence_idx: sIdx + pageStartIdx }});
                  }};
                  sSpan.appendChild(btn_undo);
              }}
              
              s.words.forEach((w, wIdx) => {{
                  let wSpan = document.createElement("span");
                  wSpan.textContent = w.word + " ";
                  
                  let classes = ["word"];
                  if (w.isLowPpl) classes.push("low-ppl");
                  if (w.isCliche) classes.push("cliche");
                  if (w.isMask) classes.push("mask");
                  wSpan.className = classes.join(" ");
                  
                  wSpan.addEventListener("click", (e) => {{
                      if(w.isLowPpl || w.isCliche) {{
                          e.preventDefault();
                          e.stopPropagation();
                          showContextMenu(e.clientX, e.clientY, sIdx + pageStartIdx, wIdx, w.word);
                      }}
                  }});
                  sSpan.appendChild(wSpan);
              }});
              
              if (s.ppl !== undefined && s.aiScore !== undefined) {{
                  let badge = document.createElement("span");
                  badge.className = "metrics-badge";
                  badge.textContent = `AI: ${{s.aiScore}}% | PPL: ${{s.ppl}}`;
                  sSpan.appendChild(badge);
              }}
              
              container.appendChild(sSpan);
          }});
          
          pendingTexts.forEach(text => {{
              let pSpan = document.createElement("span");
              pSpan.className = "sentence pending";
              pSpan.textContent = text;
              container.appendChild(pSpan);
          }});
          
          if(synonymsPayload) {{
              currentSIdx = synonymsPayload.s_idx;
              currentWIdx = synonymsPayload.w_idx;
              const menu = document.getElementById("context-menu");
              if (synonymsPayload.x !== undefined && synonymsPayload.y !== undefined) {{
                  menu.style.left = synonymsPayload.x + "px";
                  menu.style.top = synonymsPayload.y + "px";
              }}
              renderSynonymsMenu(synonymsPayload.syns_scores);
          }}
          
          // Request resize after rendering
          setTimeout(() => {{ requestResize(50); }}, 50);
      }}
  
      function showContextMenu(x, y, sIdx, wIdx, word) {{
          const menu = document.getElementById("context-menu");
          currentSIdx = sIdx;
          currentWIdx = wIdx;
          
          menu.style.left = x + "px";
          menu.style.top = y + "px";
          menu.style.display = "block";
          
          if (!synonymsPayload || synonymsPayload.s_idx !== sIdx || synonymsPayload.w_idx !== wIdx) {{
              menu.innerHTML = "<div class='loader'>Calcolo sinonimi in GPU...</div>";
              sendToGradio({{ action: "get_synonyms", sentence_idx: sIdx, word_idx: wIdx, word: word, x: x, y: y }});
          }} else {{
              renderSynonymsMenu(synonymsPayload.syns_scores);
          }}
          
          requestResize(150);
      }}
  
      function renderSynonymsMenu(syns) {{
          const menu = document.getElementById("context-menu");
          menu.innerHTML = "";
          if (syns.length === 0) {{
              menu.innerHTML = "<div class='loader'>Nessun sinonimo trovato</div>";
              return;
          }}
          syns.forEach(syn => {{
              let item = document.createElement("div");
              item.className = "menu-item";
              item.innerHTML = `${{syn.word}} <span class='score'>PPL: ${{syn.score.toFixed(1)}}</span>`;
              item.onclick = () => {{
                  menu.style.display = "none";
                  sendToGradio({{ 
                      action: "replace_word", 
                      sentence_idx: currentSIdx, 
                      word_idx: currentWIdx, 
                      new_word: syn.word 
                  }});
              }};
              menu.appendChild(item);
          }});
          menu.style.display = "block";
          requestResize(150);
      }}
  
      document.addEventListener("click", () => {{
          document.getElementById("context-menu").style.display = "none";
      }});
      
      window.addEventListener("resize", () => {{ requestResize(50); }});
      renderText();
    </script>
    </body>
    </html>
    """
    
    escaped_html = html.escape(inner_html)
    return f'<iframe srcdoc="{escaped_html}" width="100%" style="min-height: 800px; border: none; overflow: hidden;" scrolling="yes"></iframe>'

def parse_input_text(file_obj, raw_text):
    text = ""
    is_latex = False
    latex_registry = {}
    
    if file_obj is not None:
        if isinstance(file_obj, str):
            filepath = file_obj
        else:
            filepath = file_obj.name
        filename = os.path.basename(filepath).lower()
        content = open(filepath, "rb").read()
            
        if filename.endswith(".tex"):
            try:
                decoded = content.decode("utf-8")
            except:
                decoded = content.decode("latin-1")
            text, latex_registry = mask_latex(decoded)
            is_latex = True
        elif filename.endswith(".docx"):
            doc = Document(BytesIO(content))
            text = "\n".join([p.text for p in doc.paragraphs])
        else:
            try:
                text = content.decode("utf-8")
            except:
                text = content.decode("latin-1")
    elif raw_text:
        text = raw_text
        if r"\documentclass" in text or r"\begin{document}" in text:
            text, latex_registry = mask_latex(text)
            is_latex = True
            
    return text, is_latex, latex_registry

def init_analysis(file_obj, raw_text):
    print("=" * 50)
    print("[DEBUG] init_analysis CHIAMATO")
    try:
        text, is_latex, latex_registry = parse_input_text(file_obj, raw_text)
        print(f"[DEBUG] Testo estratto, lunghezza: {len(text)}")
        text = apply_algorithmic_rules(text)
        sentences = split_into_sentences(text)
        print(f"[DEBUG] Frasi trovate: {len(sentences)}")
        
        valid_sentences = []
        for s in sentences:
            if is_latex and latex_registry:
                clean_s = s
                for mask in latex_registry.keys():
                    clean_s = clean_s.replace(mask, "")
                if not any(c.isalpha() for c in clean_s):
                    continue
            valid_sentences.append(s)
            
        if not valid_sentences:
            return [], [], gr.update(visible=False), {}, False, "", "<div style='color: red; padding: 20px;'>Nessun testo valido trovato.</div>"
            
        processed_sentences = [None] * len(valid_sentences)
        import math
        total_pages = max(1, math.ceil(len(valid_sentences) / PAGE_SIZE))
        
        return valid_sentences, processed_sentences, gr.update(maximum=total_pages, value=1, visible=True), latex_registry, is_latex, text, "<div style='color: #94a3b8; padding: 20px;'>Avvio analisi pagina 1...</div>"
        
    except Exception as e:
        import traceback
        err = traceback.format_exc()
        print(f"[DEBUG] ERRORE: {err}")
        return [], [], gr.update(visible=False), {}, False, "", f"<div style='color: red; padding: 20px;'><b>ERRORE:</b><br><pre>{err}</pre></div>"

@spaces.GPU(duration=120)
def analyze_page(page_num, valid_sentences, processed_sentences, latex_registry, is_latex):
    if not valid_sentences:
        yield processed_sentences, "<div style='color: red;'>Nessun testo da analizzare.</div>"
        return
        
    page_num = max(1, int(page_num))
    start_idx = (page_num - 1) * PAGE_SIZE
    end_idx = min(start_idx + PAGE_SIZE, len(valid_sentences))
    
    html = build_html(processed_sentences, valid_sentences=valid_sentences, current_page=page_num, page_size=PAGE_SIZE)
    yield processed_sentences, html
    
    for i in range(start_idx, end_idx):
        if processed_sentences[i] is None:
            s = valid_sentences[i]
            print(f"[DEBUG] Elaboro frase {i+1}/{len(valid_sentences)}: {s[:50]}...")
            s_data = process_sentence(s, latex_registry, is_latex)
            s_data["original_text"] = s
            processed_sentences[i] = s_data
            
            yield processed_sentences, build_html(processed_sentences, valid_sentences=valid_sentences, current_page=page_num, page_size=PAGE_SIZE)
            
    print("[DEBUG] PAGINA COMPLETATA CON SUCCESSO")

@spaces.GPU(duration=60)
def handle_ui_action(payload_str, processed_sentences, latex_reg, is_latex, valid_sentences, current_page, current_language):
    print("=" * 50)
    print(f"[DEBUG] handle_ui_action CHIAMATO con payload: {payload_str}")
    if not payload_str:
        print("[DEBUG] Payload vuoto, ritorno default")
        return gr.update(), gr.update(), gr.update()
        
    try:
        action_data = json.loads(payload_str)
        action = action_data.get("action")
        print(f"[DEBUG] Azione decodificata: {action}")
        
        if action == "get_synonyms":
            s_idx = action_data["sentence_idx"]
            w_idx = action_data["word_idx"]
            word = action_data["word"]
            print(f"[DEBUG] get_synonyms per parola '{word}' a frase {s_idx}, parola {w_idx}")
            context_sentence = processed_sentences[s_idx]["text"]
            
            menu_x = action_data.get("x", 0)
            menu_y = action_data.get("y", 0)
            
            syns = get_offline_synonyms(word, context_sentence, current_language)
            print(f"[DEBUG] Sinonimi da WordNet/Mistral: {syns}")
            
            syns_scores = calculate_synonym_scores(processed_sentences[s_idx], w_idx, syns)
            print(f"[DEBUG] Score calcolati: {syns_scores}")
            
            synonyms_payload = {
                "s_idx": s_idx,
                "w_idx": w_idx,
                "syns_scores": syns_scores,
                "x": menu_x,
                "y": menu_y
            }
            return processed_sentences, build_html(processed_sentences, valid_sentences=valid_sentences, synonyms_payload=synonyms_payload, current_page=current_page), ""
            
        elif action == "replace_word":
            s_idx = action_data["sentence_idx"]
            w_idx = action_data["word_idx"]
            new_word = action_data["new_word"]
            
            new_word = new_word.replace("**", "").replace("*", "").replace("`", "")
            print(f"[DEBUG] replace_word a frase {s_idx}, parola {w_idx} con '{new_word}'")
            
            s = processed_sentences[s_idx]
            old_word = s["words"][w_idx]["word"]
            new_text = s["text"].replace(old_word, new_word, 1)
            
            original_text = s.get("original_text", s["text"])
            reprocessed = process_sentence(new_text, latex_reg, is_latex)
            reprocessed["original_text"] = original_text
            processed_sentences[s_idx] = reprocessed
            return processed_sentences, build_html(processed_sentences, valid_sentences=valid_sentences, current_page=current_page), ""
            
        elif action == "rewrite_sentence":
            s_idx = action_data["sentence_idx"]
            old_text = action_data["text"]
            print(f"[DEBUG] rewrite_sentence a frase {s_idx}: '{old_text}'")
            
            s_data = processed_sentences[s_idx]
            original_text = s_data.get("original_text", old_text)
            
            context_before = ""
            context_after = ""
            for i in range(s_idx - 1, -1, -1):
                if processed_sentences[i] is not None:
                    context_before = processed_sentences[i]["text"]
                    break
            for i in range(s_idx + 1, len(processed_sentences)):
                if processed_sentences[i] is not None:
                    context_after = processed_sentences[i]["text"]
                    break
            
            best_reprocessed = None
            # Loop fino a 3 volte cercando un punteggio AI < 30%
            for attempt in range(3):
                temp = 0.7 + (attempt * 0.2)
                new_text = rewrite_with_mistral(old_text, sentence_data=s_data, temperature=temp, lang=current_language, context_before=context_before, context_after=context_after)
                new_text = new_text.replace("**", "").replace("*", "").replace("`", "")
                print(f"[DEBUG] Tentativo {attempt+1} - Riscritto da Mistral: '{new_text}'")
                
                reprocessed = process_sentence(new_text, latex_reg, is_latex)
                
                if best_reprocessed is None or reprocessed["aiScore"] < best_reprocessed["aiScore"]:
                    best_reprocessed = reprocessed
                    
                if best_reprocessed["aiScore"] <= 30:
                    break
                    
            best_reprocessed["original_text"] = original_text
            processed_sentences[s_idx] = best_reprocessed
            return processed_sentences, build_html(processed_sentences, valid_sentences=valid_sentences, current_page=current_page), ""
            
        elif action == "undo":
            s_idx = action_data["sentence_idx"]
            print(f"[DEBUG] undo a frase {s_idx}")
            
            s_data = processed_sentences[s_idx]
            original_text = s_data.get("original_text", s_data["text"])
            
            reprocessed = process_sentence(original_text, latex_reg, is_latex)
            reprocessed["original_text"] = original_text
            processed_sentences[s_idx] = reprocessed
            return processed_sentences, build_html(processed_sentences, valid_sentences=valid_sentences, current_page=current_page), ""
            
        elif action == "manual_edit":
            s_idx = action_data["sentence_idx"]
            new_text = action_data["text"]
            print(f"[DEBUG] manual_edit a frase {s_idx}: '{new_text}'")
            
            s_data = processed_sentences[s_idx]
            original_text = s_data.get("original_text", s_data["text"])
            
            reprocessed = process_sentence(new_text, latex_reg, is_latex)
            reprocessed["original_text"] = original_text
            processed_sentences[s_idx] = reprocessed
            return processed_sentences, build_html(processed_sentences, valid_sentences=valid_sentences, current_page=current_page), ""
            
    except Exception as e:
        import traceback
        err = traceback.format_exc()
        print(f"[DEBUG] ERRORE in handle_ui_action: {err}")
        return processed_sentences, f"<div style='color: red;'><b>ERRORE:</b><pre>{err}</pre></div>", ""
        
    return processed_sentences, build_html(processed_sentences, valid_sentences=valid_sentences, current_page=current_page), ""

def export_doc(processed_sentences, latex_reg, is_latex, original_text_masked):
    if not processed_sentences: return None
    
    # Ricostruiamo il testo mantenendo gli spazi e ritorni a capo originali
    final_text = original_text_masked
    for s in processed_sentences:
        if s is not None and s.get("original_text") and s.get("original_text") != s.get("text"):
            # Sostituiamo solo la prima occorrenza per evitare di sovrascrivere frasi identiche due volte
            final_text = final_text.replace(s["original_text"], s["text"], 1)
            
    if is_latex and latex_reg:
        final_text = unmask_latex(final_text, latex_reg)
    
    if is_latex:
        path = "/tmp/humanized.tex"
        with open(path, "w", encoding="utf-8") as f:
            f.write(final_text)
        return path
    else:
        path = "/tmp/humanized.docx"
        doc = Document()
        doc.add_paragraph(final_text)
        doc.save(path)
        return path


css = """
body { background-color: #0f172a !important; color: white !important; }
.gradio-container { max-width: 95% !important; width: 95% !important; }
#action_payload { display: none !important; }
"""

head_js = """
<script>
  window.addEventListener("message", (event) => {
      if (!event.data) return;
      
      if (event.data.type === "gradio_action") {
          const textbox = document.querySelector('#action_payload textarea') || document.querySelector('#action_payload input');
          if (textbox) {
              const nativeInputValueSetter = Object.getOwnPropertyDescriptor(window.HTMLTextAreaElement.prototype, "value")?.set
                  || Object.getOwnPropertyDescriptor(window.HTMLInputElement.prototype, "value")?.set;
              if (nativeInputValueSetter) {
                  nativeInputValueSetter.call(textbox, JSON.stringify(event.data.payload));
              } else {
                  textbox.value = JSON.stringify(event.data.payload);
              }
              textbox.dispatchEvent(new Event('input', { bubbles: true }));
              textbox.dispatchEvent(new Event('change', { bubbles: true }));
          }
      } else if (event.data.type === "resize_iframe") {
          const iframe = document.querySelector('#output_html iframe');
          if (iframe) {
              iframe.style.height = event.data.height + 'px';
          }
      }
  });
</script>
"""

with gr.Blocks(css=css, head=head_js, theme=gr.themes.Default(primary_hue="blue", neutral_hue="slate")) as app:
    gr.Markdown("# 🖋️ HumanHumanHuman-nizer (Gradio ZeroGPU)")
    
    state_sentences = gr.State([])
    state_valid_sentences = gr.State([])
    state_latex_reg = gr.State({})
    state_is_latex = gr.State(False)
    state_original_text = gr.State("")
    
    with gr.Row():
        with gr.Column(scale=1):
            file_input = gr.File(label="Carica file (.txt, .docx, .tex)", file_types=[".txt", ".docx", ".tex"])
            text_input = gr.Textbox(label="O incolla il testo qui", lines=10)
            state_language = gr.Radio(["Italiano", "English"], value="Italiano", label="Lingua Documento")
            analyze_btn = gr.Button("Analizza Testo 🚀", variant="primary")
            export_btn = gr.DownloadButton("Scarica Documento Humanized 📥")
            
        with gr.Column(scale=2):
            page_slider = gr.Slider(minimum=1, maximum=1, step=1, label="Pagina", visible=False)
            output_html = gr.HTML("<div style='color: #94a3b8; padding: 20px;'>L'analisi apparirà qui...</div>", elem_id="output_html")
            
    action_payload = gr.Textbox(elem_id="action_payload")
    
    # Click analizza -> Init -> Analyze Page 1
    analyze_btn.click(
        fn=init_analysis,
        inputs=[file_input, text_input],
        outputs=[state_valid_sentences, state_sentences, page_slider, state_latex_reg, state_is_latex, state_original_text, output_html]
    ).then(
        fn=analyze_page,
        inputs=[page_slider, state_valid_sentences, state_sentences, state_latex_reg, state_is_latex],
        outputs=[state_sentences, output_html]
    )
    
    # Cambio pagina -> Analyze Page
    page_slider.change(
        fn=analyze_page,
        inputs=[page_slider, state_valid_sentences, state_sentences, state_latex_reg, state_is_latex],
        outputs=[state_sentences, output_html]
    )
    
    action_payload.change(
        fn=handle_ui_action,
        inputs=[action_payload, state_sentences, state_latex_reg, state_is_latex, state_valid_sentences, page_slider, state_language],
        outputs=[state_sentences, output_html, action_payload]
    )
    
    export_btn.click(
        fn=export_doc,
        inputs=[state_sentences, state_latex_reg, state_is_latex, state_original_text],
        outputs=[export_btn]
    )

if __name__ == "__main__":
    app.launch(share=True, server_name="0.0.0.0", server_port=7860)
